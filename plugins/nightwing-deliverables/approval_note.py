"""Generates an MRPL-style approval note as a real .docx file.

This is a *starter* template — a generic industrial approval-note layout
(reference block, findings, recommendation, sign-off table), not MRPL's
actual letterhead or house format. Swap the constants below and the
``_build_header``/``_build_signoff`` layout for the real template before
using this for an actual demo or document.

Deliberately deterministic (plain python-docx calls, not LLM-generated
markup) — reliability matters more than flexibility for a document with a
Word file as a real approval-workflow artifact.
"""

from __future__ import annotations

import os
from datetime import date as _date
from pathlib import Path
from typing import Any, Dict, List, Optional

COMPANY_NAME = "MANGALORE REFINERY AND PETROCHEMICALS LIMITED"
COMPANY_SUBTITLE = "(A Schedule 'A' Government of India Enterprise)"


def _next_ref_no(hermes_home: Path) -> str:
    """Auto-incrementing ref number, persisted under HERMES_HOME so repeat
    runs in a session don't collide. Format: NW/AN/<year>/<seq>.
    """
    year = _date.today().year
    counter_path = hermes_home / "nightwing-deliverables" / "approval_note_seq.txt"
    counter_path.parent.mkdir(parents=True, exist_ok=True)
    seq = 1
    try:
        raw = counter_path.read_text(encoding="utf-8").strip()
        stored_year, stored_seq = raw.split(":")
        if int(stored_year) == year:
            seq = int(stored_seq) + 1
    except Exception:
        pass
    counter_path.write_text(f"{year}:{seq}", encoding="utf-8")
    return f"NW/AN/{year}/{seq:03d}"


def _build_header(doc, ref_no: str, date_str: str, department: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(COMPANY_NAME)
    run.bold = True
    run.font.size = _pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(COMPANY_SUBTITLE).italic = True

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrun = heading.add_run("APPROVAL NOTE")
    hrun.bold = True
    hrun.font.size = _pt(13)
    hrun.underline = True

    doc.add_paragraph()  # spacer

    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Light Grid Accent 1"
    meta.cell(0, 0).text = "Reference No."
    meta.cell(0, 1).text = ref_no
    meta.cell(1, 0).text = "Date"
    meta.cell(1, 1).text = date_str
    if department:
        row = meta.add_row()
        row.cells[0].text = "Department"
        row.cells[1].text = department
    for row in meta.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()


def _pt(size: int):
    from docx.shared import Pt

    return Pt(size)


def _build_signoff(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ["Prepared By", "Reviewed By", "Approved By"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(3):
        cell = table.cell(1, i)
        cell.text = "\n\n\n(Name / Signature / Date)"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_approval_note(
    output_path: str,
    subject: str,
    findings: List[str],
    recommendation: str,
    hermes_home: Optional[str] = None,
    ref_no: Optional[str] = None,
    department: str = "",
    prepared_by: str = "",
    source_document: str = "",
    date_str: Optional[str] = None,
) -> Dict[str, Any]:
    """Generate an approval-note .docx file. Returns a result dict, never raises
    past its own boundary — the caller (tool handler) reports errors as text.
    """
    try:
        import docx
    except ImportError:
        return {
            "success": False,
            "error": (
                "python-docx is not installed. This is a pre-baked dependency "
                "for the air-gapped build — if it's missing, the setup step "
                "was skipped. Install with: pip install python-docx"
            ),
        }

    if not subject or not subject.strip():
        return {"success": False, "error": "subject is required"}
    if not findings:
        return {"success": False, "error": "findings must be a non-empty list"}
    if not recommendation or not recommendation.strip():
        return {"success": False, "error": "recommendation is required"}

    home = Path(hermes_home) if hermes_home else Path.home() / ".hermes"
    resolved_ref = ref_no or _next_ref_no(home)
    resolved_date = date_str or _date.today().strftime("%d %B %Y")

    doc = docx.Document()
    _build_header(doc, resolved_ref, resolved_date, department)

    doc.add_paragraph().add_run("Subject: ").bold = True
    doc.paragraphs[-1].add_run(subject)

    if source_document:
        p = doc.add_paragraph()
        p.add_run("Source Document: ").bold = True
        p.add_run(source_document)

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.add_run("Findings").bold = True
    h.runs[0].font.size = _pt(12)
    for item in findings:
        doc.add_paragraph(str(item), style="List Number")

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    h2.add_run("Recommendation").bold = True
    h2.runs[0].font.size = _pt(12)
    doc.add_paragraph(recommendation)

    if prepared_by:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Prepared by: ").bold = True
        p.add_run(prepared_by)

    _build_signoff(doc)

    out = Path(output_path).expanduser()
    if not out.is_absolute():
        out = Path.cwd() / out
    if out.suffix.lower() != ".docx":
        out = out.with_suffix(".docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    return {
        "success": True,
        "path": str(out),
        "ref_no": resolved_ref,
        "date": resolved_date,
        "message": f"Approval note saved to {out} (ref {resolved_ref}).",
    }
